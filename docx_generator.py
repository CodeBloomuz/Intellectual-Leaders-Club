import shutil
from docx import Document
from config import TEMPLATE_PATH


def _replace_in_paragraph(para, replacements: dict):
    """
    Paragraf ichidagi matnni almashtiradi.
    Run-lar bo'lingan bo'lsa ham ishlaydi — birinchi run'ga birlashtiradi.
    """
    full_text = "".join(run.text for run in para.runs)
    if not any(key in full_text for key in replacements):
        return

    new_text = full_text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def generate_docx(data: dict, familya_initial: str) -> str:
    fish = data["full_name"]
    guruh = data["guruh"]

    safe_name = fish.replace(" ", "_")
    output_path = f"ariza_{safe_name}.docx"

    shutil.copy(TEMPLATE_PATH, output_path)
    doc = Document(output_path)

    # Templateda mavjud matnlarni almashtiramiz
    # Muhim: uzunroq stringlar avval almashtirilishi kerak
    replacements = {
        "Iqtisodiyot va axborot texnologiyalari fakulteti": data["fakultet"],
        "Iqtisodiyot yo'nalishi": data["yonalish"],
        "Toshtemirov Dilmurod Xasanovich": fish,
        "IQT-4-23-guruh": f"{guruh}-guruh",
        "IQT-4-23": guruh,
        "4-23-guruh": f"{guruh}-guruh",
        "4-23": guruh,
        "D.X.Toshtemirov": familya_initial,
    }

    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    doc.save(output_path)
    return output_path


def make_familya_initial(fish: str) -> str:
    """
    'Toshtemirov Dilmurod Xasanovich' → 'D.X.Toshtemirov'
    """
    parts = fish.strip().split()
    if len(parts) >= 3:
        return f"{parts[1][0]}.{parts[2][0]}.{parts[0]}"
    elif len(parts) == 2:
        return f"{parts[1][0]}.{parts[0]}"
    return fish
